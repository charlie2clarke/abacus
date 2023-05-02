import docx
from models.document import Document


class DocumentDao:
    ''' Data access object for interacting with word documents '''

    def __init__(self, file_path):
        self.document = Document(file_path)
        if self.document.doc == None:
            return

        self.initialise_document()

    @property
    def document_loaded(self):
        return self.document.doc is not None

    def initialise_document(self):
        self.document.tables = self.document.doc.tables
        self.document.paras = list(self.set_paragraphs())
        self.document.words = list(self.set_words())

    def set_paragraphs(self):
        counter = 0

        while counter < len(self.document.doc.paragraphs):
            paragraph = self.document.doc.paragraphs[counter]
            counter += 1

            if paragraph.text.startswith('Appendix') or paragraph.text.startswith('Appendices'):
                break

            if not paragraph.style.name.startswith('Heading') and not paragraph.style.name.startswith('Title') and not paragraph.style.name.startswith('Subtitle') and \
                    paragraph.text != '' and not paragraph.text.startswith('Figure'):
                yield paragraph

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text != '':
                            yield paragraph

    def set_words(self):
        for paragraph in self.document.paras:
            paragraph = paragraph.text
            words = paragraph.split()
            for word in words:
                if word != ('.' or '–'):
                    yield word

    def print_word_count(self):
        print(f'Words: {len(self.document.words)}')
